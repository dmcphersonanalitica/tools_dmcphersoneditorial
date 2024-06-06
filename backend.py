import os
import streamlit as st
import whisper
import tempfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

quality_model = {
    '1': 'tiny',
    '2': 'base',
    '3': 'small',
    '4': 'medium',
    '5': 'large'
    }


@st.cache_data
@st.cache_resource
def load_model(quality):
    try:
        model = whisper.load_model(quality_model[str(quality)])
        return model
    except Exception as e:
        print(e)


def audio_transcription(model, audio_file, format):
    options = whisper.DecodingOptions(fp16=False)
    result = model.transcribe(audio_file)

    if format == 'Timeline text':
        full_result = ''
        segments = result["segments"]
        for seg in segments:
            full_result += str(seg['start']) + '   --->   ' + str(seg['end']) + '\n\n' + seg['text'] + '\n\n'

        return full_result

    return result["text"]


def path_file(file):
    temp_dir = tempfile.mkdtemp()
    path = os.path.join(temp_dir, file.name)
    with open(path, "wb") as f:
        f.write(file.getvalue())
    return path.replace('\\', '\\\\')


def save_to_word(filename, text, format):
    document = Document()

    header = document.add_heading('DMcPherson Editorial Tools\n')
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(text)
    
    if format == 'Timeline text':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save(filename.name + '.docx')
    return document
